from tkinter import *
import tkinter as tk
from tkinter import ttk, filedialog
import os
import json
import random
from docx.shared import Pt
from docx import Document

class GenerateMCQs(tk.Frame):
    def __init__(self, parentRoot):
        super().__init__(parentRoot)
        self.pack(padx=20, pady=20)

        # Title Label
        self.label = ttk.Label(self, text="Generate MCQs", font=("Arial", 16, "bold"))
        self.label.pack(pady=(0, 10))

        # Frame for MCQ Input and Options
        self.GenerateMCQsFrame = ttk.Frame(self)
        self.GenerateMCQsFrame.pack(fill=tk.BOTH, expand=True)

        # Frame for Inputs
        self.GenerateMCQsInputsFrame = ttk.Frame(self.GenerateMCQsFrame)
        self.GenerateMCQsInputsFrame.pack(pady=10)

        # File Selection Label and Combobox
        self.GenerateMCQsLabel = ttk.Label(self.GenerateMCQsInputsFrame, text="Select a file:")
        self.GenerateMCQsLabel.grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)

        options = self.GetAllMCQFiles()
        self.GenerateMCQsFileSelect = ttk.Combobox(self.GenerateMCQsInputsFrame, values=options, state="readonly", width=30)
        self.GenerateMCQsFileSelect.grid(row=0, column=1, padx=10, pady=5)

        # Generate Button with custom style
        self.GenerateMCQsButton = tk.Button(
            self.GenerateMCQsInputsFrame,
            text="Generate MCQs",
            font=("Segoe UI", 10, "bold"),
            bg="#3498db",  # Blue background
            fg="white",  # White text
            activebackground="#2980b9",  # Darker blue on hover
            activeforeground="white",  # White text on hover
            relief="flat",  # Flat button design
            cursor="hand2",  # Hand cursor on hover
            command=self.DisplayQuestionsInTextBox
        )
        self.GenerateMCQsButton.grid(row=0, column=2, padx=10, pady=5)


        # Text Area for Displaying Questions
        self.textArea = tk.Text(self.GenerateMCQsFrame, wrap=WORD, height=20, width=60, font=("Courier", 10))
        self.textArea.pack(pady=10)

        self.saveButton = tk.Button(
        self.GenerateMCQsFrame,
        text="Save to Word",
        font=("Segoe UI", 10, "bold"),
        bg="#28a745",  # Green background
        fg="white",  # White text
        activebackground="#218838",  # Darker green for hover
        activeforeground="white",  # White text on hover
        relief="flat",  # Flat button design
        cursor="hand2",  # Hand cursor on hover
        command=self.SaveQuestionsToWord
        )
        self.saveButton.pack(pady=(5, 10))

        
        self.saveButton = tk.Button(
        self.GenerateMCQsFrame,
        text="Create Question Bank",
        font=("Segoe UI", 10, "bold"),
        bg="#28a745",  # Green background
        fg="white",  # White text
        activebackground="#218838",  # Darker green for hover
        activeforeground="white",  # White text on hover
        relief="flat",  # Flat button design
        cursor="hand2",  # Hand cursor on hover
        command=self.MakeQuestionBank
        )
        self.saveButton.pack(pady=(5, 10))

        # Update Message Label
        self.updateMessage = ttk.Label(self.GenerateMCQsFrame, text="", foreground="green")
        self.updateMessage.pack()

        # Initialize Questions List
        self.questions = []

    def GetTheDataInTheSelectedFile(self):
        selectedFile = self.GenerateMCQsFileSelect.get().strip()
        current_folder_path = os.path.dirname(os.path.abspath(__file__))
        data_folder_path = os.path.join(current_folder_path, '..', 'data', 'MCQs')

        if not os.path.exists(data_folder_path):
            self.updateMessage["text"] = "Data folder does not exist!"
            return {}

        filePath = os.path.join(data_folder_path, f'{selectedFile}.json')
        if not os.path.exists(filePath):
            self.updateMessage["text"] = "File does not exist!"
            return {}

        try:
            with open(filePath, 'r', encoding='utf-8') as jsonFile:
                data = json.load(jsonFile)
            return data
        except json.JSONDecodeError:
            self.updateMessage["text"] = "Invalid JSON file!"
            return {}
        except Exception as e:
            self.updateMessage["text"] = f"Error reading file: {str(e)}"
            return {}

    def GenerateMCQPaper(self, data):
        try:
            units = list(data.keys())
            self.selected_questions = []

            for unit in units:
                questions = data.get(unit, [])
                if len(questions) < 4:
                    self.updateMessage["text"] = f"Not enough questions in {unit}. Required: 4."
                    return []
                self.selected_questions.extend(random.sample(questions, 4))

            random.shuffle(self.selected_questions)
            return self.selected_questions
        except Exception as e:
            self.updateMessage["text"] = f"Error generating MCQ paper: {str(e)}"
            return []

    def DisplayQuestionsInTextBox(self):
        self.data = self.GetTheDataInTheSelectedFile()
        if not self.data:
            return

        self.questions = self.GenerateMCQPaper(self.data)
        if not self.questions:
            return

        self.textArea.delete(1.0, END)
        for i, question in enumerate(self.questions, 1):
            self.textArea.insert(END, f"Q{i}. {question['question']}\n")
            for j in range(1, 5):
                self.textArea.insert(END, f"  Option {j}: {question[f'option{j}']}\n")
            self.textArea.insert(END, "\n")

    def SaveQuestionsToWord(self):
        try:
            if isinstance(self.data, str):
                self.data = json.loads(self.data)

            if not isinstance(self.data, dict):
                self.updateMessage["text"] = "Invalid questions data format!"
                return

            # Create a Word document
            document = Document()

            # Add title with custom font size
            title = document.add_heading('Multiple Choice Questions (MCQ) Paper', level=1)
            title_run = title.runs[0]
            title_run.font.size = Pt(14)  # Title font size

            # Add sets of questions
            for set in range(4):
                if set > 0:  # Add a page break before each new set (except the first set)
                    document.add_page_break()

                shuffled_questions = self.questions.copy()  # Create a copy to avoid modifying the original list
                random.shuffle(shuffled_questions)
                set_heading = document.add_heading(f'SET-{set + 1}', level=2)
                set_heading_run = set_heading.runs[0]
                set_heading_run.font.size = Pt(12)  # Font size for set headings

                # Add questions
                for i, question in enumerate(shuffled_questions, 1):
                    question_text = question.get('question', '')
                    option1 = question.get('option1', '')
                    option2 = question.get('option2', '')
                    option3 = question.get('option3', '')
                    option4 = question.get('option4', '')

                    # Add question with reduced line spacing and smaller font
                    question_paragraph = document.add_paragraph(f"{i}. {question_text}\t\t[   ]", style='Normal')
                    question_paragraph.paragraph_format.line_spacing = Pt(12)  # Reduce line spacing
                    for run in question_paragraph.runs:
                        run.font.size = Pt(10)  # Adjust font size for question

                    # Add options
                    options_paragraph = document.add_paragraph(
                        f"1. {option1}\t2. {option2}\t3. {option3}\t4. {option4}\n", style='Normal'
                    )
                    options_paragraph.paragraph_format.line_spacing = Pt(12)  # Reduce line spacing
                    for run in options_paragraph.runs:
                        run.font.size = Pt(10)  # Adjust font size for options

            # Open save dialog
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
            if save_path:
                document.save(save_path)
                self.updateMessage["text"] = f"MCQ Paper saved successfully to {save_path}"
            else:
                self.updateMessage["text"] = "Save operation canceled."

        except Exception as e:
            self.updateMessage["text"] = f"Error saving Word document: {e}"
    

    def SaveQuestionBankToWord(self, QuestionBankQuestions):
        try:
            # Create a Word document
            document = Document()

            # Add title
            title = document.add_heading('Question Bank', level=1)
            title_run = title.runs[0]
            title_run.font.size = Pt(14)  # Adjust title font size

            # Add questions
            for i, question in enumerate(QuestionBankQuestions, 1):
                question_text = question.get('question', '')
                option1 = question.get('option1', '')
                option2 = question.get('option2', '')
                option3 = question.get('option3', '')
                option4 = question.get('option4', '')

                # Add question
                question_paragraph = document.add_paragraph(f"{i}. {question_text}\t\t[   ]")
                question_paragraph.paragraph_format.line_spacing = Pt(12)  # Reduce line spacing
                for run in question_paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size

                # Add options
                options_paragraph = document.add_paragraph(
                    f"1. {option1}\t2. {option2}\t3. {option3}\t4. {option4}\n"
                )
                options_paragraph.paragraph_format.line_spacing = Pt(12)  # Reduce line spacing
                for run in options_paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size

            # Open save dialog
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
            if save_path:
                document.save(save_path)
                self.updateMessage["text"] = f"Question Bank saved successfully to {save_path}"
            else:
                self.updateMessage["text"] = "Save operation canceled."

        except Exception as e:
            self.updateMessage["text"] = f"Error saving Question Bank to Word: {e}"

    
    def MakeQuestionBank(self):
        # This will take the already selcted questions data and adds new questions to that data 
        # Use set and select random questions until the set length becomes 30-40 bits
        # then save that file into the question bank
        # This question bank will be sent to students as a study material

        questionBankData = self.selected_questions.copy()

        TotalQuestionsData = self.data.copy()

        QuestionsSeenInQuestionBankData = set(q['question'] for q in questionBankData)

        new_question_Bank = questionBankData.copy()

        all_questions = []
        for unit_questions in TotalQuestionsData.values():
            all_questions.extend(unit_questions)

        random.shuffle(all_questions)

        while len(new_question_Bank) < 30:
            for question in all_questions:
                if question['question'] not in QuestionsSeenInQuestionBankData:
                    new_question_Bank.append(question)
                    QuestionsSeenInQuestionBankData.add(question['question'])

                if len(new_question_Bank) >= 30:
                    break        
        
        self.SaveQuestionBankToWord(new_question_Bank)

        print(new_question_Bank)
        print(len(new_question_Bank))
            

    def GetAllMCQFiles(self):
        current_folder_path = os.path.dirname(os.path.abspath(__file__))
        folder_path = os.path.join(current_folder_path, '..', 'data', 'MCQs')
        if not os.path.exists(folder_path):
            return []

        return [f[:-5] for f in os.listdir(folder_path) if f.endswith(".json")]
