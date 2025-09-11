# GanerateQuestions.py
# -*- coding: utf-8 -*-
"""
GenerateQuestions tkinter Frame.
Implements:
- Load questions JSON
- Normalize question pools
- Generate Mid-I (Unit1 vs Unit2, 3 numbered questions)
- Generate Mid-II (Units 3-5 pairings, 3 numbered questions)
- Generate SEM (as before)
- Preview and Save to Word
"""

import json
import os
import random
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime

# Optional dependency: python-docx for saving .docx
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ---------- Helper utilities ----------

def normalize_unit_keys(raw_unit_dict):
    """
    Convert keys like '2 Marks', '2', 'Two' to integer keys 2, 5, 10.
    Returns dict {2: [...], 5: [...], 10: [...]} (missing keys absent).
    """
    normalized = {}
    for k, v in raw_unit_dict.items():
        if not isinstance(k, str):
            keystr = str(k).lower()
        else:
            keystr = k.lower().replace(" ", "").replace("-", "")
        if keystr.startswith("2"):
            normalized.setdefault(2, []).extend(v if isinstance(v, list) else [v])
        elif keystr.startswith("5"):
            normalized.setdefault(5, []).extend(v if isinstance(v, list) else [v])
        elif keystr.startswith("10"):
            normalized.setdefault(10, []).extend(v if isinstance(v, list) else [v])
        else:
            # attempt fuzzy matches
            if "2" in keystr or "two" in keystr:
                normalized.setdefault(2, []).extend(v if isinstance(v, list) else [v])
            elif "5" in keystr or "five" in keystr:
                normalized.setdefault(5, []).extend(v if isinstance(v, list) else [v])
            elif "10" in keystr or "ten" in keystr:
                normalized.setdefault(10, []).extend(v if isinstance(v, list) else [v])
            else:
                # ignore unknown key
                pass
    # normalize entries
    for mark in list(normalized.keys()):
        cleaned = []
        for entry in normalized[mark]:
            if isinstance(entry, dict) and "question" in entry:
                cleaned.append({"question": entry.get("question", "").strip(),
                                "bt": entry.get("bt", "").strip() if entry.get("bt") else ""})
            elif isinstance(entry, str):
                cleaned.append({"question": entry.strip(), "bt": ""})
            else:
                continue
        normalized[mark] = cleaned
    return normalized

def normalize_questions_json(raw):
    normalized_all = {}
    for unit_name, unit_data in raw.items():
        normalized_all[unit_name] = normalize_unit_keys(unit_data if isinstance(unit_data, dict) else {})
    return normalized_all

def pick_random_question_from_pool(pool):
    if not pool:
        return None
    return random.choice(pool)

# ---------- GenerateQuestions Frame ----------

class GenerateQuestions(tk.Frame):
    def __init__(self, parentRoot):
        super().__init__(parentRoot)
        self.parent = parentRoot
        self.pack(fill="both", expand=True)

        # State
        self.questions_data = {}  # normalized data
        self.current_filepath = None
        self.generated_text = ""
        self.seed_value = None

        # UI components
        self._build_ui()

        # Try default file in cwd
        default_path = os.path.join(os.getcwd(), "questions_data.json")
        if os.path.exists(default_path):
            self.load_questions_file(default_path)

    # ---------- UI building ----------
    def _build_ui(self):
        top = tk.Frame(self, bg="#0b1220", pady=8)
        top.pack(fill="x")

        lbl = tk.Label(top, text="Question Paper Generator", font=("Helvetica", 14, "bold"),
                       fg="white", bg="#0b1220")
        lbl.pack(side="left", padx=12)

        btn_load = tk.Button(top, text="Load file", command=self._on_load_file,
                             bg="#10b981", fg="white", activebackground="#059669", padx=8)
        btn_load.pack(side="right", padx=8)
        btn_save = tk.Button(top, text="Save to Word", command=self._on_save_word,
                             bg="#3b82f6", fg="white", activebackground="#2563eb", padx=8)
        btn_save.pack(side="right", padx=8)

        row = tk.Frame(self, pady=6)
        row.pack(fill="x", padx=12)

        tk.Label(row, text="Mode:", font=("Helvetica", 10)).pack(side="left")
        self.mode_var = tk.StringVar(value="Mid-I")
        mode_combo = ttk.Combobox(row, textvariable=self.mode_var, values=["Mid-I", "Mid-II", "SEM"], width=10)
        mode_combo.pack(side="left", padx=(6, 12))
        mode_combo.state(["readonly"])

        tk.Label(row, text="Seed (optional):", font=("Helvetica", 10)).pack(side="left")
        self.seed_entry = tk.Entry(row, width=8)
        self.seed_entry.pack(side="left", padx=(6, 12))

        self.bt_var = tk.BooleanVar(value=True)
        bt_chk = tk.Checkbutton(row, text="Show BT tags", variable=self.bt_var)
        bt_chk.pack(side="left")

        gen_btn = tk.Button(row, text="Generate", command=self._on_generate,
                            bg="#ef4444", fg="white", activebackground="#dc2626", padx=12)
        gen_btn.pack(side="right", padx=6)

        regen_btn = tk.Button(row, text="Regenerate", command=self._on_regenerate,
                              bg="#f59e0b", fg="white", activebackground="#d97706", padx=12)
        regen_btn.pack(side="right", padx=6)

        mid = tk.Frame(self)
        mid.pack(fill="both", expand=True, padx=12, pady=8)

        left_panel = tk.Frame(mid)
        left_panel.pack(side="left", fill="y")

        tk.Label(left_panel, text="Loaded file:", font=("Helvetica", 9, "bold")).pack(anchor="w")
        self.lbl_file = tk.Label(left_panel, text="(none)", fg="#374151")
        self.lbl_file.pack(anchor="w")

        tk.Label(left_panel, text="Units preview:", font=("Helvetica", 9, "bold"), pady=6).pack(anchor="w")
        self.units_listbox = tk.Listbox(left_panel, height=10, width=36)
        self.units_listbox.pack()

        right_panel = tk.Frame(mid)
        right_panel.pack(side="right", fill="both", expand=True)

        tk.Label(right_panel, text="Generated Paper Preview", font=("Helvetica", 10, "bold")).pack(anchor="w")
        self.preview = tk.Text(right_panel, wrap="word", state="disabled", padx=8, pady=8)
        self.preview.pack(fill="both", expand=True)

        self.preview.tag_config("header", font=("Helvetica", 12, "bold"), foreground="#111827")
        self.preview.tag_config("unit", font=("Helvetica", 10, "bold"), foreground="#0ea5a4")
        self.preview.tag_config("question", font=("Helvetica", 10), foreground="#0f172a")
        self.preview.tag_config("bt", font=("Helvetica", 9, "italic"), foreground="#6b7280")
        self.preview.tag_config("or", font=("Helvetica", 10, "bold"), foreground="#b91c1c")

        bottom = tk.Frame(self, pady=6)
        bottom.pack(fill="x")
        self.status_lbl = tk.Label(bottom, text="Ready.", anchor="w", fg="#374151")
        self.status_lbl.pack(fill="x", padx=12)

    # ---------- File loading ----------
    def _on_load_file(self):
        path = filedialog.askopenfilename(initialdir=".", title="Select questions JSON",
                                          filetypes=[("JSON files", "*.json"), ("All files", "*.*")])
        if path:
            self.load_questions_file(path)

    def load_questions_file(self, path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                raw = json.load(f)
        except Exception as e:
            messagebox.showerror("Load error", f"Failed to load JSON:\n{e}")
            self.status_lbl.config(text="Failed to load file.")
            return

        normalized = normalize_questions_json(raw)
        if not normalized:
            messagebox.showerror("Format error", "Loaded JSON contains no valid units/questions.")
            self.status_lbl.config(text="Invalid file.")
            return

        self.questions_data = normalized
        self.current_filepath = path
        self.lbl_file.config(text=os.path.basename(path))
        self.units_listbox.delete(0, tk.END)
        for unit in self.questions_data.keys():
            unit_info = self.questions_data[unit]
            s = f"{unit}  |  2m: {len(unit_info.get(2,[]))}  5m: {len(unit_info.get(5,[]))}  10m: {len(unit_info.get(10,[]))}"
            self.units_listbox.insert(tk.END, s)
        self.status_lbl.config(text=f"Loaded {len(self.questions_data)} units from file.")
        self._clear_preview()

    # ---------- Generators ----------
    def _on_generate(self):
        seed_text = self.seed_entry.get().strip()
        if seed_text:
            try:
                seed_int = int(seed_text)
                random.seed(seed_int)
                self.seed_value = seed_int
            except Exception:
                messagebox.showwarning("Seed", "Seed must be an integer. Ignoring.")
                self.seed_value = None
                random.seed()
        else:
            self.seed_value = None
            random.seed()

        mode = self.mode_var.get()
        if not self.questions_data:
            messagebox.showwarning("No data", "No questions file loaded. Use 'Load file' or place questions_data.json in app folder.")
            return

        try:
            if mode == "SEM":
                out = self._generate_sem()
            elif mode == "Mid-I":
                out = self._generate_mid(mid_type=1)
            elif mode == "Mid-II":
                out = self._generate_mid(mid_type=2)
            else:
                raise ValueError("Unknown mode")
        except ValueError as e:
            messagebox.showerror("Validation error", str(e))
            self.status_lbl.config(text="Generation failed.")
            return

        self.generated_text = out
        self._render_preview(out)
        self.status_lbl.config(text=f"Generated paper ({mode}).")

    def _on_regenerate(self):
        if self.seed_value is not None:
            random.seed(self.seed_value)
        else:
            random.seed()
        self._on_generate()

    def _generate_mid(self, mid_type=1):
        """
        Updated behavior:
        - Mid-I (mid_type=1): use only Unit1 and Unit2 (first two units). Produce 3 numbered questions;
            for each question display Unit1 alternative OR Unit2 alternative; each alternative = 10 marks.
        - Mid-II (mid_type=2): use Unit3, Unit4, Unit5 (units at indices 2,3,4). Produce 3 numbered questions with pairings:
            Q1: Unit3 vs Unit4
            Q2: Unit3 vs Unit5
            Q3: Unit4 vs Unit5
        """
        units = list(self.questions_data.keys())
        n_units = len(units)

        def prepare_local_pools(unit_name):
            u = self.questions_data.get(unit_name, {})
            # shallow copies so we can pop without affecting original in case we want to reuse original for fallback
            return {
                10: u.get(10, []).copy(),
                5: u.get(5, []).copy(),
                2: u.get(2, []).copy()
            }

        def select_10_or_5pair(unit_name, local_pools, orig_pools):
            """
            Try to pick a fresh 10m from local_pools; if not possible, pick two 5m from local_pools.
            If local pools exhausted, fallback to random choice from orig_pools and return note='(reused)'.
            Returns dict with keys: type (10 or '5+5'), text (string), bt (string), reused (bool)
            """
            reused = False
            lp = local_pools[unit_name]
            op = orig_pools[unit_name]
            # try 10-mark fresh
            if lp[10]:
                q10 = lp[10].pop(random.randrange(len(lp[10])))
                return {"type": 10, "text": f"{unit_name} - {q10['question']}", "bt": q10.get("bt", ""), "reused": False}
            # try two fresh 5-marks
            if len(lp[5]) >= 2:
                a = lp[5].pop(random.randrange(len(lp[5])))
                b = lp[5].pop(random.randrange(len(lp[5])))
                return {"type": "5+5", "text": f"{unit_name} - (a) {a['question']}  (b) {b['question']}", "bt": f"{a.get('bt','')} / {b.get('bt','')}", "reused": False}
            if len(lp[5]) == 1:
                a = lp[5].pop(0)
                # duplicate it (only one available)
                return {"type": "5+5", "text": f"{unit_name} - (a) {a['question']}  (b) {a['question']}  (NOTE: only one 5m available)", "bt": a.get("bt",""), "reused": False}

            # Fallback (reuse from original pools)
            reused = True
            if op.get(10):
                q10 = random.choice(op[10])
                return {"type": 10, "text": f"{unit_name} - {q10['question']}  (reused)", "bt": q10.get("bt", ""), "reused": True}
            if op.get(5):
                # try distinct if possible
                if len(op[5]) >= 2:
                    a, b = random.sample(op[5], 2)
                    return {"type": "5+5", "text": f"{unit_name} - (a) {a['question']}  (b) {b['question']}  (reused)", "bt": f"{a.get('bt','')} / {b.get('bt','')}", "reused": True}
                else:
                    a = op[5][0]
                    return {"type": "5+5", "text": f"{unit_name} - (a) {a['question']}  (b) {a['question']}  (reused)", "bt": a.get("bt",""), "reused": True}
            # nothing at all
            return None

        # Mid-I rules
        if mid_type == 1:
            if n_units < 2:
                raise ValueError("Mid-I requires at least 2 units (Unit1 and Unit2) in the questions file.")
            unitA = units[0]
            unitB = units[1]
            # prepare pools
            local_pools = {unitA: prepare_local_pools(unitA), unitB: prepare_local_pools(unitB)}
            orig_pools = {unitA: {"10": self.questions_data[unitA].get(10, []).copy(), "5": self.questions_data[unitA].get(5, []).copy()},
                          unitB: {"10": self.questions_data[unitB].get(10, []).copy(), "5": self.questions_data[unitB].get(5, []).copy()}}
            slots = []
            any_reused = False
            # generate 3 questions (Q1..Q3)
            for i in range(3):
                altA = select_10_or_5pair(unitA, local_pools, orig_pools)
                altB = select_10_or_5pair(unitB, local_pools, orig_pools)
                if altA is None or altB is None:
                    raise ValueError(f"Unable to find any suitable questions for Mid-I from '{unitA}' and/or '{unitB}'.")
                if altA.get("reused") or altB.get("reused"):
                    any_reused = True
                slots.append((altA, altB))

            # format output
            lines = []
            lines.append("MID TERM EXAM\n")
            for idx, (a, b) in enumerate(slots, start=1):
                # show unit A alt then OR then unit B alt
                if a["type"] == 10:
                    lines.append(f"Q{idx}. {a['text']}  [10 marks]")
                else:
                    lines.append(f"Q{idx}. {a['text']}  [10 marks]")
                lines.append("     OR")
                if b["type"] == 10:
                    lines.append(f"     {b['text']}  [10 marks]")
                else:
                    lines.append(f"     {b['text']}  [10 marks]")
                if self.bt_var.get():
                    lines.append(f"    BT: {a.get('bt','')}   OR   {b.get('bt','')}")
                lines.append("")
            if any_reused:
                lines.append("(Note: some alternatives were reused due to limited question pool.)")
            return "\n".join(lines)

        # Mid-II rules
        else:
            # require units 3,4,5 present
            if n_units < 5:
                raise ValueError("Mid-II requires at least 5 units in the questions file (so Unit3, Unit4, Unit5 are present).")
            u3, u4, u5 = units[2], units[3], units[4]
            # pairings: (3 vs 4), (3 vs 5), (4 vs 5)
            pairings = [(u3, u4), (u3, u5), (u4, u5)]
            local_pools = {u3: prepare_local_pools(u3), u4: prepare_local_pools(u4), u5: prepare_local_pools(u5)}
            orig_pools = {
                u3: {"10": self.questions_data[u3].get(10, []).copy(), "5": self.questions_data[u3].get(5, []).copy()},
                u4: {"10": self.questions_data[u4].get(10, []).copy(), "5": self.questions_data[u4].get(5, []).copy()},
                u5: {"10": self.questions_data[u5].get(10, []).copy(), "5": self.questions_data[u5].get(5, []).copy()},
            }
            slots = []
            any_reused = False
            for (left, right) in pairings:
                altL = select_10_or_5pair(left, local_pools, orig_pools)
                altR = select_10_or_5pair(right, local_pools, orig_pools)
                if altL is None or altR is None:
                    raise ValueError(f"Unable to find suitable questions for Mid-II pair: '{left}' vs '{right}'.")
                if altL.get("reused") or altR.get("reused"):
                    any_reused = True
                slots.append((altL, altR))
            # format output
            lines = []
            lines.append("MID TERM EXAM - (Mid II)\n")
            for idx, (a, b) in enumerate(slots, start=1):
                lines.append(f"Q{idx}. {a['text']}  [10 marks]")
                lines.append("     OR")
                lines.append(f"     {b['text']}  [10 marks]")
                if self.bt_var.get():
                    lines.append(f"    BT: {a.get('bt','')}   OR   {b.get('bt','')}")
                lines.append("")
            if any_reused:
                lines.append("(Note: some alternatives were reused due to limited question pool.)")
            return "\n".join(lines)

    # ---------- SEM generator (unchanged) ----------
    def _generate_sem(self):
        units = list(self.questions_data.keys())
        if len(units) != 5:
            raise ValueError("SEM generation requires exactly 5 units in the questions file (one per unit).")

        # Part A
        partA = []
        for unit in units:
            pool2 = self.questions_data[unit].get(2, [])
            if not pool2:
                raise ValueError(f"SEM Part A requires at least one 2-mark question in '{unit}'.")
            q2 = pick_random_question_from_pool(pool2)
            partA.append({"unit": unit, "text": q2["question"], "bt": q2.get("bt", "")})

        # Part B
        partB = []
        for unit in units:
            pools = self.questions_data[unit]
            alternatives = []
            if pools.get(10):
                q10 = pick_random_question_from_pool(pools[10])
                alternatives.append({"type": 10, "text": q10["question"], "bt": q10.get("bt", "")})
            if len(pools.get(5, [])) >= 2:
                a, b = random.sample(pools[5], 2)
                alternatives.append({"type": "5+5", "text": (a["question"], b["question"]), "bt": f"{a.get('bt','')} / {b.get('bt','')}"})
            elif len(pools.get(5, [])) == 1:
                a = pools[5][0]
                alternatives.append({"type": "5+5", "text": (a["question"], a["question"]), "bt": a.get("bt","") + " (dup)"})
            if not alternatives:
                raise ValueError(f"Unit '{unit}' lacks both 10m and 5m questions required for SEM Part B.")
            if len(alternatives) == 1:
                alternatives.append({"type": alternatives[0]["type"], "text": alternatives[0]["text"], "bt": alternatives[0].get("bt","") + " (only)"} )
            a_alt, b_alt = random.sample(alternatives, 2)
            partB.append({"unit": unit, "alt1": a_alt, "alt2": b_alt})

        lines = []
        lines.append("SEMESTER EXAM\n")
        lines.append("PART A: (Answer all - 2 Marks each)\n")
        for i, q in enumerate(partA, start=1):
            lines.append(f"{i}. ({q['unit']}) {q['text']}  [2 marks]")
            if self.bt_var.get() and q.get("bt"):
                lines.append(f"    BT: {q.get('bt')}")
        lines.append("\nPART B: (For each unit, answer ONE of the following - 10 marks each)\n")

        qnum = 1
        for pb in partB:
            lines.append(f"Unit: {pb['unit']}")
            alt1 = pb['alt1']
            alt2 = pb['alt2']
            if alt1["type"] == 10:
                lines.append(f"{qnum}. (a) {alt1['text']}  [10 marks]")
            else:
                a1, b1 = alt1["text"]
                lines.append(f"{qnum}. (a) (i) {a1}  (ii) {b1}  [5 + 5 = 10 marks]")
            lines.append("     OR")
            if alt2["type"] == 10:
                lines.append(f"     (b) {alt2['text']}  [10 marks]")
            else:
                a2, b2 = alt2["text"]
                lines.append(f"     (b) (i) {a2}  (ii) {b2}  [5 + 5 = 10 marks]")
            if self.bt_var.get():
                lines.append(f"    BT: {alt1.get('bt','')}   OR   {alt2.get('bt','')}")
            lines.append("")
            qnum += 1

        return "\n".join(lines)

    # ---------- Preview rendering ----------
    def _clear_preview(self):
        self.preview.config(state="normal")
        self.preview.delete("1.0", tk.END)
        self.preview.config(state="disabled")

    def _render_preview(self, text):
        self.preview.config(state="normal")
        self.preview.delete("1.0", tk.END)

        lines = text.splitlines()
        for ln in lines:
            stripped = ln.strip()
            if not stripped:
                self.preview.insert(tk.END, "\n")
                continue
            if stripped.upper().startswith("MID TERM") or stripped.upper().startswith("SEMESTER"):
                self.preview.insert(tk.END, ln + "\n", "header")
            elif stripped.startswith("PART A") or stripped.startswith("PART B"):
                self.preview.insert(tk.END, ln + "\n", "header")
            elif stripped.startswith("Unit:"):
                self.preview.insert(tk.END, ln + "\n", "unit")
            elif stripped.strip() == "OR" or stripped.startswith("OR"):
                self.preview.insert(tk.END, ln + "\n", "or")
            elif "BT:" in stripped:
                self.preview.insert(tk.END, ln + "\n", "bt")
            else:
                self.preview.insert(tk.END, ln + "\n", "question")

        self.preview.config(state="disabled")

    # ---------- Save to Word ----------
    def _on_save_word(self):
        if not self.generated_text:
            messagebox.showwarning("No paper", "No generated paper to save. Click Generate first.")
            return
        if not DOCX_AVAILABLE:
            messagebox.showerror("Missing dependency", "python-docx is not installed. Install it to enable Word export:\n\npip install python-docx")
            return

        default_name = f"Generated_Paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=default_name,
                                            filetypes=[("Word Document", "*.docx")])
        if not path:
            return

        try:
            self._save_docx(path, self.generated_text)
            messagebox.showinfo("Saved", f"Saved Word document to:\n{path}")
            self.status_lbl.config(text=f"Saved: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Save error", f"Failed to save DOCX:\n{e}")
            self.status_lbl.config(text="Save failed.")

    def _save_docx(self, path, text):
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        lines = text.splitlines()
        for ln in lines:
            if not ln.strip():
                doc.add_paragraph("")
                continue
            s = ln.strip()
            if s.upper().startswith("MID TERM") or s.upper().startswith("SEMESTER"):
                p = doc.add_paragraph()
                run = p.add_run(s)
                run.bold = True
                run.font.size = Pt(14)
            elif s.startswith("PART A") or s.startswith("PART B"):
                p = doc.add_paragraph()
                run = p.add_run(s)
                run.bold = True
                run.font.size = Pt(12)
            elif s.startswith("Unit:"):
                p = doc.add_paragraph()
                run = p.add_run(s)
                run.bold = True
            else:
                p = doc.add_paragraph(s)
        doc.save(path)

# Minimal test
if __name__ == "__main__":
    root = tk.Tk()
    root.title("Question Generator - Demo")
    root.geometry("900x640")
    frame = GenerateQuestions(root)
    root.mainloop()
